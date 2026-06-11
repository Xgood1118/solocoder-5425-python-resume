import io
import re
import logging
from typing import Tuple, List, Dict, Optional
from .base import BaseParser

logger = logging.getLogger(__name__)

HEADER_PATTERNS = {
    'date': [
        '时间', '日期', '起止时间', '时间段', '年月', '年份',
        '入学时间', '毕业时间', '开始时间', '结束时间',
        'date', 'time', 'period', 'year', 'from', 'to'
    ],
    'school': [
        '学校', '院校', '毕业院校', '所在学校', '就读学校',
        'school', 'university', 'college', 'institute', 'academy'
    ],
    'major': [
        '专业', '所学专业', '专业名称', '主修', '专业方向',
        'major', 'specialty', 'field', 'subject'
    ],
    'degree': [
        '学位', '学历', '文凭', 'degree', 'diploma', 'education'
    ],
    'company': [
        '公司', '单位', '工作单位', '所属公司', '雇主',
        'company', 'employer', 'organization', 'firm'
    ],
    'position': [
        '职位', '岗位', '职务', '担任职务', 'title', 'position', 'post', 'role'
    ],
    'department': [
        '部门', '所属部门', '所在部门', 'department', 'dept', 'division'
    ],
    'description': [
        '描述', '工作内容', '职责', '主要工作', '业绩', '成果',
        'description', 'responsibility', 'achievement', 'duty'
    ],
}


class WordParser(BaseParser):
    def supports(self, filename: str) -> bool:
        return filename.lower().endswith(('.docx', '.doc'))

    def parse(self, file_content: bytes, filename: str) -> Tuple[str, dict]:
        metadata = {
            "parse_method": "word",
            "confidence": 1.0,
            "paragraph_count": 0,
            "table_count": 0,
        }
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            metadata["parse_method"] = "word_error"
            metadata["confidence"] = 0.0
            metadata["error"] = "python-docx not installed"
            return "", metadata

        try:
            text_parts = []
            doc = Document(io.BytesIO(file_content))

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            metadata["paragraph_count"] = len(doc.paragraphs)

            for table in doc.tables:
                metadata["table_count"] += 1
                table_text = self._parse_table(table)
                if table_text:
                    text_parts.append(table_text)

            full_text = "\n".join(text_parts).strip()
            metadata["confidence"] = 0.95

            return full_text, metadata

        except Exception as e:
            logger.error(f"Word parsing failed: {e}")
            metadata["parse_method"] = "word_error"
            metadata["confidence"] = 0.0
            metadata["error"] = str(e)
            return "", metadata

    def _parse_table(self, table) -> Optional[str]:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return None

        header_row_idx, header_map = self._detect_header(rows_data)

        if header_map:
            return self._parse_structured_table(rows_data, header_row_idx, header_map)
        else:
            return self._parse_unstructured_table(rows_data)

    def _detect_header(self, rows_data: List[List[str]]) -> Tuple[int, Dict[int, str]]:
        for row_idx in range(min(3, len(rows_data))):
            row = rows_data[row_idx]
            header_map: Dict[int, str] = {}
            for col_idx, cell in enumerate(row):
                if not cell:
                    continue
                cell_lower = cell.lower().strip()
                for field, patterns in HEADER_PATTERNS.items():
                    if any(p in cell_lower for p in patterns):
                        header_map[col_idx] = field
                        break
            if len(header_map) >= 2:
                return row_idx, header_map
        return -1, {}

    def _parse_structured_table(self, rows_data: List[List[str]], header_row_idx: int, header_map: Dict[int, str]) -> str:
        output_lines = []

        has_education = any(f in header_map.values() for f in ['school', 'major', 'degree'])
        has_work = any(f in header_map.values() for f in ['company', 'position', 'department'])

        if has_education:
            output_lines.append("教育经历")
        elif has_work:
            output_lines.append("工作经历")

        for row_idx in range(header_row_idx + 1, len(rows_data)):
            row = rows_data[row_idx]
            if not any(row):
                continue

            fields: Dict[str, str] = {}
            for col_idx, field_type in header_map.items():
                if col_idx < len(row) and row[col_idx]:
                    fields[field_type] = row[col_idx].strip()

            if not fields:
                continue

            date_part = fields.get('date', '')
            school_part = fields.get('school', '')
            company_part = fields.get('company', '')
            major_part = fields.get('major', '')
            degree_part = fields.get('degree', '')
            position_part = fields.get('position', '')
            department_part = fields.get('department', '')
            desc_part = fields.get('description', '')

            if has_education:
                line_parts = []
                if date_part:
                    line_parts.append(date_part)
                if school_part:
                    line_parts.append(school_part)
                if major_part:
                    line_parts.append(major_part)
                if degree_part:
                    line_parts.append(degree_part)
                if line_parts:
                    output_lines.append("  " + "  ".join(line_parts))
                if desc_part:
                    output_lines.append("    " + desc_part)
            elif has_work:
                line_parts = []
                if date_part:
                    line_parts.append(date_part)
                if company_part:
                    line_parts.append(company_part)
                if department_part:
                    line_parts.append(department_part)
                if position_part:
                    line_parts.append(position_part)
                if line_parts:
                    output_lines.append("  " + "  ".join(line_parts))
                if desc_part:
                    output_lines.append("    " + desc_part)
            else:
                line_parts = []
                for col_idx, field_type in header_map.items():
                    if col_idx < len(row) and row[col_idx]:
                        line_parts.append(row[col_idx].strip())
                if line_parts:
                    output_lines.append("  " + "  ".join(line_parts))

        return "\n".join(output_lines) if len(output_lines) > 1 else None

    def _parse_unstructured_table(self, rows_data: List[List[str]]) -> str:
        output_lines = []
        for row in rows_data:
            non_empty_cells = [cell for cell in row if cell.strip()]
            if not non_empty_cells:
                continue

            row_has_date = any(re.search(r'(20\d{2}|19\d{2})', cell) for cell in non_empty_cells)
            row_has_school = any(any(kw in cell for kw in ['大学', '学院', '学校', 'University', 'College', 'Institute']) for cell in non_empty_cells)
            row_has_company = any(any(kw in cell for kw in ['公司', '集团', '科技', '有限', '有限公司', '股份']) for cell in non_empty_cells)

            if len(non_empty_cells) >= 2 and (row_has_date or row_has_school or row_has_company):
                output_lines.append("  " + "  ".join(non_empty_cells))
            else:
                output_lines.append(" | ".join(non_empty_cells))

        return "\n".join(output_lines) if output_lines else None
